"""Parser for yes/no keyed Word tests.

The parser reads table-based questions, validates the yes/no scoring keys,
and turns the document into a structured test model.
"""

from docx import Document

from school_form_app.models import ParsedTest, Question, AnswerOption


DEFAULT_YES_SCORE_QUESTIONS = {
    1, 3, 5, 6, 8, 10, 11, 13, 15, 19
}

DEFAULT_NO_SCORE_QUESTIONS = {
    2, 4, 7, 9, 12, 14, 16, 17, 18, 20
}


def clean_text(text: str) -> str:
    return " ".join(text.strip().split())


def extract_title(document: Document) -> str:
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if text:
            return text

    return "Untitled Yes/No test"


def extract_instructions(document: Document) -> str | None:
    instruction_lines = []
    inside_instruction = False

    stop_words = [
        "ключ",
        "подсчет",
        "подсчёт",
        "обработка",
        "интерпретация",
    ]

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if not text:
            continue

        text_lower = text.lower()

        if inside_instruction:
            if any(text_lower.startswith(word) for word in stop_words):
                break

        if text_lower.startswith("инструкция"):
            inside_instruction = True

        if inside_instruction:
            instruction_lines.append(text)

    if not instruction_lines:
        return None

    return "\n".join(instruction_lines)


def parse_question_numbers(raw_text: str) -> set[int]:
    """
    Converts text like:
        "1,3,5,6"
        "1, 3, 5-8"
        "1;3;5"
    into:
        {1, 3, 5, 6, 7, 8}
    """
    result = set()

    cleaned = (
        raw_text
        .replace(";", ",")
        .replace("–", "-")
        .replace("—", "-")
    )

    parts = cleaned.split(",")

    for part in parts:
        part = part.strip()

        if not part:
            continue

        if "-" in part:
            start_text, end_text = part.split("-", 1)

            start = int(start_text.strip())
            end = int(end_text.strip())

            if start > end:
                raise ValueError(
                    f"Wrong range in key: {part}. "
                    f"Start number cannot be bigger than end number."
                )

            for number in range(start, end + 1):
                result.add(number)

        else:
            result.add(int(part))

    return result


def validate_question_key(
    question_number: int,
    yes_score_questions: set[int],
    no_score_questions: set[int],
) -> tuple[int, int]:
    is_yes_scored = question_number in yes_score_questions
    is_no_scored = question_number in no_score_questions

    if is_yes_scored and is_no_scored:
        raise ValueError(
            f"Вопрос {question_number} находится и в ДА=1, и в НЕТ=1.\n"
            f"Один вопрос не может быть в двух ключах одновременно."
        )

    if not is_yes_scored and not is_no_scored:
        raise ValueError(
            f"Вопрос {question_number} не добавлен в ключ обработки.\n\n"
            f"Добавь его либо в список ДА=1, либо в список НЕТ=1."
        )

    yes_score = 1 if is_yes_scored else 0
    no_score = 1 if is_no_scored else 0

    return yes_score, no_score


def parse_yes_no_table_questions(
    document: Document,
    yes_score_questions: set[int],
    no_score_questions: set[int],
) -> list[Question]:
    # Each Word table row represents one question, and the first two columns
    # contain the question number and its statement.
    questions = []
    seen_question_numbers = set()

    for table in document.tables:
        for row in table.rows:
            cells = [
                clean_text(cell.text)
                for cell in row.cells
            ]

            if len(cells) < 2:
                continue

            number_text = cells[0]
            question_text = cells[1]

            if not number_text.isdigit():
                continue

            question_number = int(number_text)

            if question_number in seen_question_numbers:
                continue

            if not question_text:
                continue

            yes_score, no_score = validate_question_key(
                question_number=question_number,
                yes_score_questions=yes_score_questions,
                no_score_questions=no_score_questions,
            )

            question = Question(
                number=question_number,
                title=question_text,
                options=[
                    AnswerOption(text="ДА", score=yes_score),
                    AnswerOption(text="НЕТ", score=no_score),
                ],
            )

            questions.append(question)
            seen_question_numbers.add(question_number)

    if not questions:
        raise ValueError(
            "No Yes/No table questions found in the Word file.\n\n"
            "Проверь, что вопросы находятся в таблице Word, "
            "где первый столбец — номер, второй столбец — утверждение."
        )

    questions.sort(key=lambda question: question.number)

    return questions


def parse_yes_no_keyed_docx(
    path: str,
    yes_score_questions: set[int] | None = None,
    no_score_questions: set[int] | None = None,
) -> ParsedTest:
    if yes_score_questions is None:
        yes_score_questions = DEFAULT_YES_SCORE_QUESTIONS

    if no_score_questions is None:
        no_score_questions = DEFAULT_NO_SCORE_QUESTIONS

    overlap = yes_score_questions.intersection(no_score_questions)

    if overlap:
        raise ValueError(
            f"Эти вопросы есть одновременно в ДА=1 и НЕТ=1: "
            f"{sorted(overlap)}"
        )

    document = Document(path)

    title = extract_title(document)
    instructions = extract_instructions(document)

    questions = parse_yes_no_table_questions(
        document=document,
        yes_score_questions=yes_score_questions,
        no_score_questions=no_score_questions,
    )

    return ParsedTest(
        title=title,
        instructions=instructions,
        questions=questions,
    )