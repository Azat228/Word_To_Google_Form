"""
Parser for Likert-style tests.

A Likert-style test is a test where every question uses the same answer options.

Examples:

1. Ferguson loneliness test:
    Often = 3
    Sometimes = 2
    Rarely = 1
    Never = 0

2. Burnout test:
    Never = 0
    Rarely = 1
    Sometimes = 2
    Often = 3
    Very often = 4

The important idea:
    The Word file contains the questions.
    The JSON config contains the answer options and scores.

So this parser combines:

    Word questions + config answer options

and creates a ParsedTest object that the rest of the app can use.
"""

from logging import config
import re

from docx import Document

from school_form_app.models import ParsedTest, Question, AnswerOption


# This regular expression finds numbered questions.
#
# It matches lines like:
#
#   1. I feel tired after work
#   2) I feel lonely
#
# Explanation:
#   ^\s*        -> allow spaces at the beginning
#   (\d+)       -> capture the question number
#   [\.\)]      -> allow either "." or ")" after the number
#   \s*         -> allow spaces after "." or ")"
#   (.+)        -> capture the question text
QUESTION_RE = re.compile(r"^\s*(\d+)[\.\)]\s*(.+)")


def clean_text(text: str) -> str:
    """
    Clean text copied from Word.

    Word files often contain:
    - extra spaces,
    - line breaks,
    - tabs,
    - invisible formatting characters.

    This function converts text like:

        "   Hello     world   "

    into:

        "Hello world"
    """

    return " ".join(text.strip().split())


def extract_title(document: Document) -> str:
    """
    Extract the first non-empty paragraph as the test title.

    Example:

        МЕТОДИКА ДИАГНОСТИКИ УРОВНЯ СУБЪЕКТИВНОГО ОЩУЩЕНИЯ ОДИНОЧЕСТВА

    If no title is found, we return a fallback title.
    """

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if text:
            return text

    return "Untitled Likert test"


def extract_instructions(document: Document) -> str | None:
    """
    Extract instructions from the Word file.

    The idea:
        Take paragraphs from the beginning of the document
        until we reach the first numbered question or result-processing section.

    This is not perfect for every possible Word file,
    but it works well for simple psychological test blanks.

    Example:

        Instruction. Read each statement and choose one answer...
    """

    instruction_lines = []

    # These words usually mean that the instruction part is over.
    stop_words = [
        "подсчет",
        "подсчёт",
        "обработка",
        "интерпретация",
        "итого",
        "ключ",
    ]

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if not text:
            continue

        text_lower = text.lower()

        # If we reached a numbered question, stop collecting instructions.
        if QUESTION_RE.match(text):
            break

        # If we reached result-processing text, stop collecting instructions.
        if any(text_lower.startswith(word) for word in stop_words):
            break

        instruction_lines.append(text)

    if not instruction_lines:
        return None

    return "\n".join(instruction_lines)


def make_options_from_config(config: dict) -> list[AnswerOption]:
    """
    Create answer options from JSON config.

    Example config:

        "answer_options": [
            {"text": "Часто", "score": 3},
            {"text": "Иногда", "score": 2},
            {"text": "Редко", "score": 1},
            {"text": "Никогда", "score": 0}
        ]

    This function converts each dictionary into an AnswerOption object.
    """

    options = []

    for option_data in config.get("answer_options", []):
        option = AnswerOption(
            text=str(option_data["text"]),
            score=int(option_data["score"]),
        )

        options.append(option)

    if not options:
        raise ValueError(
            "Config has no answer_options. "
            "Every Likert test must define answer options."
        )

    return options


def clone_options(options: list[AnswerOption]) -> list[AnswerOption]:
    """
    Create a fresh copy of answer options for each question.

    Why not reuse the same AnswerOption objects?

    Because later we may modify one question's options,
    and we do not want that modification to affect all other questions.

    This is safer:

        Question 1 gets its own option objects.
        Question 2 gets its own option objects.
        Question 3 gets its own option objects.
    """

    cloned = []

    for option in options:
        cloned.append(
            AnswerOption(
                text=option.text,
                score=option.score,
            )
        )

    return cloned


def should_stop_reading(text: str) -> bool:
    """
    Decide whether we should stop reading questions.

    Many Word files contain text after the questions, for example:

        Подсчет баллов
        Обработка результатов
        Интерпретация
        ИТОГО

    These are not questions.
    When we see them, we stop parsing.
    """

    text_lower = text.lower()

    stop_words = [
        "подсчет",
        "подсчёт",
        "обработка",
        "интерпретация",
        "итого",
        "ключ",
    ]

    return any(
        text_lower.startswith(word)
        for word in stop_words
    )


def parse_questions_from_paragraphs(
    document: Document,
    config: dict,
) -> list[Question]:
    """
    Parse questions from normal Word paragraphs.

    This is used for files like the burnout test.

    Example Word structure:

        1. I feel emotionally exhausted after work
        2. At the end of the working day, I have no energy
        3. In the morning, it is hard for me to go to work

    Each numbered paragraph becomes one Question object.

    The answer options come from the JSON config.
    """

    questions = []
    base_options = make_options_from_config(config)

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)

        if not text:
            continue

        if should_stop_reading(text):
            break

        match = QUESTION_RE.match(text)

        if not match:
            continue

        question_number = int(match.group(1))
        question_title = match.group(2).strip()

        if not question_title:
            raise ValueError(
                f"Question {question_number} has empty text."
            )

        question = Question(
            number=question_number,
            title=question_title,
            options=clone_options(base_options),
        )

        questions.append(question)

    return questions

def normalize_for_compare(text: str) -> str:
    """
    Normalize text for comparison.

    We use this only to compare cell text with answer option labels.

    Example:
        "  Часто " -> "часто"
        "Никогда" -> "никогда"
    """

    return clean_text(text).lower().replace("ё", "е")


def get_answer_option_texts(config: dict) -> set[str]:
    """
    Get all answer option texts from config.

    Example:
        ["Часто", "Иногда", "Редко", "Никогда"]

    We use this to avoid accidentally taking an answer option as question text.
    """

    option_texts = set()

    for option_data in config.get("answer_options", []):
        option_text = normalize_for_compare(
            str(option_data.get("text", ""))
        )

        if option_text:
            option_texts.add(option_text)

    return option_texts


def extract_question_text_from_table_cells(
    cells: list[str],
    config: dict,
) -> str:
    """
    Try to find the question text inside a table row.

    Normal expected structure:

        cells[0] = question number
        cells[1] = question text

    But some Word files are messy.
    Sometimes cells[1] is empty, and the question text is shifted
    to another cell.

    This function tries:
        1. use cells[1] if it exists and is not empty;
        2. otherwise search other cells;
        3. ignore answer option labels like "Часто", "Иногда", etc.
    """

    if len(cells) >= 2 and cells[1]:
        return cells[1]

    answer_option_texts = get_answer_option_texts(config)

    # Start from cell 1 because cell 0 is the question number.
    for cell_text in cells[1:]:
        normalized_cell_text = normalize_for_compare(cell_text)

        if not normalized_cell_text:
            continue

        # Do not use answer options as question text.
        if normalized_cell_text in answer_option_texts:
            continue

        # Ignore very short technical cells.
        if len(normalized_cell_text) <= 1:
            continue

        return cell_text

    return ""

def parse_questions_from_tables(
    document: Document,
    config: dict,
) -> list[Question]:
    """
    Parse questions from Word tables.

    This is used for files like the Ferguson loneliness test.

    Expected table structure:

        Column 1: question number
        Column 2: question text
        Column 3+: answer columns

    Some Word files may contain broken rows, for example:

        12 | empty | empty | empty | empty

    In that case, we do not immediately crash.
    We skip the broken row and continue parsing.

    Later, parse_likert_docx(...) will check expected_question_count.
    If the final number of parsed questions is wrong, the user will get
    a clearer error message.
    """

    questions = []
    seen_numbers = set()
    skipped_empty_question_numbers = []
    base_options = make_options_from_config(config)

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [
                clean_text(cell.text)
                for cell in row.cells
            ]

            if len(cells) < 2:
                continue

            number_text = cells[0]

            if not number_text.isdigit():
                continue

            question_number = int(number_text)

            if question_number in seen_numbers:
                continue

            question_text = extract_question_text_from_table_cells(
                cells=cells,
                config=config,
            )

            if not question_text:
                skipped_empty_question_numbers.append(question_number)

                print(
                    f"Warning: skipped empty table row. "
                    f"Question number: {question_number}, "
                    f"table: {table_index}, row: {row_index}"
                )

                continue

            question = Question(
                number=question_number,
                title=question_text,
                options=clone_options(base_options),
            )

            questions.append(question)
            seen_numbers.add(question_number)

    if skipped_empty_question_numbers:
        print()
        print(
            "Warning: these question numbers had empty text and were skipped:"
        )
        print(skipped_empty_question_numbers)
        print()

    return questions

def parse_likert_docx(
    docx_path: str,
    config: dict,
) -> ParsedTest:
    """
    Main function for Likert-style Word files.

    Input:
        docx_path:
            Path to the selected Word file.

        config:
            JSON test template.

    Output:
        ParsedTest object.

    This function checks config["parser_type"]:

        "likert_table"
            -> read questions from Word tables

        "numbered_questions_with_options"
            -> read questions from normal numbered paragraphs

    After questions are parsed, the rest of the app can use the result exactly
    like any other ParsedTest.
    """

    document = Document(docx_path)

    title = extract_title(document)
    instructions = extract_instructions(document)

    parser_type = config.get("parser_type")

    if parser_type == "likert_table":
        questions = parse_questions_from_tables(
            document=document,
            config=config,
        )

    elif parser_type == "numbered_questions_with_options":
        questions = parse_questions_from_paragraphs(
            document=document,
            config=config,
        )

    else:
        raise ValueError(
            f"Unsupported Likert parser_type: {parser_type}"
        )

    if not questions:
        raise ValueError(
            "No Likert questions found in the Word file."
        )

    expected_question_count = config.get("expected_question_count")

    if expected_question_count is not None:
        expected_question_count = int(expected_question_count)

    if len(questions) != expected_question_count:
        found_numbers = [
            question.number
            for question in questions
        ]

        raise ValueError(
            f"Wrong number of questions parsed.\n\n"
            f"Expected: {expected_question_count}\n"
            f"Found: {len(questions)}\n"
            f"Found question numbers: {found_numbers}\n\n"
            f"Most likely, the Word file has an empty or broken question row. "
            f"Please check the missing question number in the Word document."
        )

    questions.sort(
        key=lambda question: question.number
    )

    return ParsedTest(
        title=title,
        instructions=instructions,
        questions=questions,
    )